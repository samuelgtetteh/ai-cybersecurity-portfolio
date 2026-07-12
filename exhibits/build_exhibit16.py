"""
Generate "Exhibit 16 - Real-Time Security Decisioning Platform.docx".

Documents the operational decisioning system built on top of the three detectors: the
RedMap backend now Records every verdict, Decides (policy rules -> alerts, weighted by
historical outcomes), Acts (severity-routed responders + audit), and produces AI triage
(RAG over the compliance corpus + a local LLM summary). Mirrors the Exhibit 14 style and,
like it, reports MEASURED results from an actual live run.

All figures are measured from a representative live run recorded 2026-07-12 (RedMap serving
the current code; the two live-target-lab event sources streaming continuously and reporting
ground truth via the feedback channel). Reproducible via notebooks 08/09 and tests/.

Run:  venv\\Scripts\\python.exe exhibits\\build_exhibit16.py
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

RUN_DATE = "July 12, 2026"
REPO = "https://github.com/samuelgtetteh/ai-cybersecurity-portfolio"

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = bold; r.italic = italic
    return p


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        run = t.rows[0].cells[j].paragraphs[0].add_run(h); run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = str(v)
    return t

# ---- title block ----
title = doc.add_paragraph()
r = title.add_run("Exhibit 16: Real-Time Security Decisioning Platform "
                  "(Record → Decide → Act, with AI-Assisted Triage)")
r.bold = True; r.font.size = Pt(15)
para("Petitioner: Samuel Gbli Tetteh", bold=True)
para("NIW Proposed Endeavor — Operationalizing the three detection models into a single "
     "real-time decisioning system: turning live anomaly detection (validated in Exhibit 14) "
     "into recorded, triaged, and actioned security decisions.", italic=True)

# 1. Purpose
doc.add_heading("1. Purpose", level=1)
para(
    "Exhibit 14 demonstrated that the three detection models operate correctly as live, "
    "containerized services. This exhibit documents the next stage I built on top of them: an "
    "operational decisioning platform that does not merely score events but keeps an "
    "authoritative record of every decision, applies policy to raise prioritized alerts, takes "
    "auditable response actions, and produces AI-assisted analyst triage. It is the difference "
    "between a detector that flags anomalies and a system that manages them. All results below "
    "are measured from an actual live run; the system, code, and reproductions are public.")

# 2. Architecture
doc.add_heading("2. Architecture", level=1)
para("A single containerized FastAPI backend (“RedMap”) hosts the pipeline. Each stage "
     "is additive and independently auditable:")
table(["Stage", "What it does", "Interface"],
      [["Record", "Persists every model verdict to a durable trail, enriched with request "
        "metadata; a decoupled feedback channel attaches ground truth after the fact, from "
        "which live precision/recall are computed.",
        "/identity/score, /ics/score, /map; /decision/verdicts, /stats, /metrics; "
        "X-Verdict-Id header; /decision/verdicts/{id}/feedback"],
       ["Decide", "A configurable policy engine reads the trail and raises durable alerts "
        "(identity burst, sustained ICS anomalies, single high-severity events), with severity "
        "weighted by each subject’s historical outcomes (chronic false positives suppressed, "
        "confirmed-bad escalated).",
        "/decision/alerts, /decision/evaluate"],
       ["Act", "Severity-routed responders fire when an alert is raised (log, ticket, webhook); "
        "every response is recorded for audit; analysts close alerts.",
        "/decision/actions, /decision/alerts/{id}/close"],
       ["AI triage", "For an alert, retrieves the most relevant compliance controls (RAG over "
        "the HIPAA corpus using the fine-tuned RegMap embedder) and writes a concise analyst "
        "triage summary with a local LLM.",
        "/decision/alerts/{id}/triage"]])

# 3. Measured live results
doc.add_heading("3. Measured Live Results", level=1)
para(f"The following are measured from a representative live run recorded {RUN_DATE}, with the "
     "backend serving the current code and the two event sources streaming continuously and "
     "reporting ground truth through the feedback channel.")

doc.add_heading("3.1 Record — coverage and live detection metrics", level=2)
para("Every verdict was recorded and automatically labelled with ground truth via the feedback "
     "channel, letting the system compute detection quality live from its own trail:")
table(["Quantity", "Value"],
      [["Verdicts recorded (and labelled)", "1,167 (100% labelled)"],
       ["By model", "identity 762 (448 flagged); ICS 405 (55 flagged)"],
       ["Non-scored requests audited", "1,172"],
       ["Live precision", "0.992"],
       ["Live recall", "0.977"],
       ["Live specificity", "0.994"],
       ["Confusion matrix (TP/FP/FN/TN)", "499 / 4 / 12 / 652"]])
para("These are computed by comparing each model’s flagged decision against the ground truth "
     "attached after the fact — the same comparison the Exhibit 14 tally performed on a small "
     "sample, now computed continuously at scale and persisted in the database.")

doc.add_heading("3.2 Decide — alerts raised", level=2)
para("Over the run the policy engine raised five alerts, correctly isolating the three attacker "
     "accounts and the sustained/severe ICS activity from the benign majority:")
table(["Rule", "Subject", "Severity"],
      [["identity_burst", "ANONYMOUS LOGON@C9999", "high"],
       ["identity_burst", "ANONYMOUS LOGON@C8123", "high"],
       ["identity_burst", "guest@DOM1", "high"],
       ["ics_sustained", "(12 sustained ICS anomalies)", "high"],
       ["high_severity", "(single extreme ICS reading)", "high"]])
para("Each burst alert recorded the subject’s historical outcomes (e.g. malicious: 192, "
     "benign: 0), which the policy used to escalate severity — a decision informed by history, "
     "not only by the current score.")

doc.add_heading("3.3 Act — responses taken (audited)", level=2)
para("For each high-severity alert the Act layer fired and recorded its responders: a log "
     "acknowledgement, a ticket stub (with a generated reference such as SEC-8 and a descriptive "
     "title), and a webhook (cleanly skipped when no webhook URL is configured). Analysts close "
     "alerts through the API; a closed alert is suppressed from re-firing for its window.")

doc.add_heading("3.4 AI triage — grounded analyst summary", level=2)
para("For an identity-burst alert, the triage layer retrieved the most relevant HIPAA provisions "
     "via the fine-tuned RegMap embedder — the top match being “Unique User Identification "
     "(assign a unique name/number for identifying and tracking user identity)”, which is "
     "directly on point for a credential-based lateral-movement alert. A local language model then "
     "produced a concise, AI-generated triage (what happened, why it matters, recommended next "
     "action) grounded in those retrieved controls — recommending, for example, review of "
     "unique-user-identification policies and automatic-logoff enforcement. The retrieval and "
     "templated summary run within the deployed container; the language-model summary runs where "
     "the local model is present.")

# 4. Feedback loop improving decisions
doc.add_heading("4. The Feedback Loop Improves the Decisions", level=1)
para("Because ground truth is captured for each decision, the system does more than react to raw "
     "scores. A subject whose flagged logins are consistently confirmed benign is treated as a "
     "chronic false positive and its alerts are suppressed, while a subject with confirmed-"
     "malicious history is escalated. In a controlled reproduction, an identical behavioural burst "
     "raised a high-severity alert for an attacker account but was suppressed for a service "
     "account with a benign history — the alert queue stays trustworthy as the system learns "
     "from analyst/operator feedback.")

# 5. Reproducibility
doc.add_heading("5. Reproducibility and Verifiability", level=1)
para("The entire platform is public and independently reproducible:")
bullet("notebooks/08_decision_layer_record.ipynb — the Record layer and feedback/metrics loop.")
bullet("notebooks/09_decision_layer_decide_act.ipynb — Decide and Act, including the "
       "outcome-weighted suppression.")
bullet("tests/test_decision_layer.py — a seven-test automated suite exercising the full "
       "Record → Decide → Act pipeline on an isolated database.")
bullet("backend/ (verdict_store.py, policy.py, actions.py, ai_triage.py, decision_api.py) — "
       "the implementation, served from the same reproducible Docker image as the detectors.")
p = doc.add_paragraph(); p.add_run(REPO).bold = True

# 6. Significance
doc.add_heading("6. Significance for the “Well Positioned” Prong", level=1)
para("This exhibit shows that I have carried the proposed endeavor beyond building and validating "
     "individual detection models to engineering the operational system around them — recording "
     "decisions durably, prioritizing them by policy and learned history, acting on them with an "
     "auditable trail, and applying AI to accelerate analyst triage. It demonstrates end-to-end "
     "capability across research, machine learning, systems engineering, and operational security, "
     "and it reuses my own published research artifacts (the fine-tuned compliance embedder) as a "
     "component of the platform. Together with Exhibits 11–15, it is direct evidence that I am "
     "well positioned to advance this endeavor independently.")

para("")
para("Date: " + RUN_DATE)
para("Prepared by: Samuel Gbli Tetteh")

out = r"c:\Users\User\ai-cybersecurity-portfolio\exhibits\Exhibit 16 Real-Time Security Decisioning Platform.docx"
doc.save(out)
print("saved:", out)
print("paragraphs:", len(doc.paragraphs), "| tables:", len(doc.tables))
