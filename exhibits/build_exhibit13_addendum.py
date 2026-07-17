"""Generate "Exhibit 13A - OT-ICS Intrusion Detector Released as a Standalone System.docx"."""
from docx import Document
from docx.shared import Pt

RUN_DATE = "July 17, 2026"
REPO = "https://github.com/samuelgtetteh/ai-cybersecurity-portfolio"

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


def para(t, bold=False, italic=False):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = bold; r.italic = italic; return p


def bullet(t):
    doc.add_paragraph(t, style="List Bullet")


def table(headers, rows):
    tb = doc.add_table(rows=1, cols=len(headers)); tb.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        tb.rows[0].cells[j].paragraphs[0].add_run(h).bold = True
    for row in rows:
        c = tb.add_row().cells
        for j, v in enumerate(row):
            c[j].text = str(v)
    return tb


t = doc.add_paragraph()
r = t.add_run("Exhibit 13A: OT/ICS Intrusion Detector — Released as a Standalone System")
r.bold = True; r.font.size = Pt(15)
para("Petitioner: Samuel Gbli Tetteh", bold=True)
para("Addendum to Exhibit 13 — documenting that the OT/ICS intrusion detector has been packaged and "
     "publicly released as a reusable, standalone tool.", italic=True)

doc.add_heading("1. Packaged and published for public use", level=1)
para("The OT/ICS Intrusion Detector (a deep autoencoder that detects cyber-physical attacks on "
     "industrial control systems from sensor telemetry, trained on normal operation only) has been "
     "released as a self-contained tool anyone can run, under the permissive Apache-2.0 license.")
table(["Channel", "Location"],
      [["Hugging Face", "huggingface.co/stetteh/otics-anomaly"],
       ["GitHub Release", "v0.1-otics (weights + scaler + metadata + inference wrapper, runs offline)"],
       ["Docker image", "ghcr.io/samuelgtetteh/otics-anomaly (serving API: POST /score, GET /example)"]])

doc.add_heading("2. What the release contains", level=1)
bullet("The trained autoencoder weights, the fitted scaler, and the metadata (input/encoding "
       "dimensions, detection threshold, and the 59 sensor feature order) used at training time.")
bullet("A scoring wrapper that reconstructs a reading and flags it when the reconstruction error "
       "exceeds the trained threshold, filling any missing sensor tags from a real normal reading.")
bullet("A FastAPI serving image (POST /score, GET /example) and a runnable example.")
bullet("A model card stating intended use and limitations — assistive monitoring; the threshold is "
       "tuned to the HAI testbed and must be re-fit for a different plant's normal data.")

doc.add_heading("3. Significance", level=1)
para("Detecting cyber-physical attacks on critical infrastructure is a national priority; releasing "
     "this detector openly — usable at no cost by water, energy, manufacturing, and healthcare "
     "facility operators — reinforces both the national importance of the endeavor and that the "
     "petitioner is well positioned to advance it. The detector is also reused within the "
     "integrated platform (Exhibits 16–19).")
p = doc.add_paragraph(); p.add_run(REPO).bold = True
para(""); para("Date: " + RUN_DATE); para("Prepared by: Samuel Gbli Tetteh")

out = r"c:\Users\User\ai-cybersecurity-portfolio\exhibits\Exhibit 13A OT-ICS Detector Standalone Release.docx"
doc.save(out); print("saved:", out, "| paras", len(doc.paragraphs), "| tables", len(doc.tables))
