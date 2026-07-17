"""
Generate "Exhibit 11A - RegMap as a Reusable, Published Model.docx" — an addendum to Exhibit 11.

Documents that the RegMap embedder has moved beyond a single research prototype: it is now (1)
packaged and released as an open, reusable model anyone can run, and (2) reused as a component
across the operational platform (control mapping in the assessment tool, and semantic interpretation
of the compliance interview). Mirrors the Exhibit 14/16/17 style.

Run:  venv\\Scripts\\python.exe exhibits\\build_exhibit11_addendum.py
"""
from docx import Document
from docx.shared import Pt

RUN_DATE = "July 16, 2026"
REPO = "https://github.com/samuelgtetteh/ai-cybersecurity-portfolio"

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = bold; r.italic = italic
    return p


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        run = t.rows[0].cells[j].paragraphs[0].add_run(h); run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = str(v)
    return t


title = doc.add_paragraph()
r = title.add_run("Exhibit 11A: RegMap as a Reusable, Published Model")
r.bold = True; r.font.size = Pt(15)
para("Petitioner: Samuel Gbli Tetteh", bold=True)
para("Addendum to Exhibit 11 (RegMap) — documenting that the RegMap compliance-mapping model has "
     "been released as an open, reusable artifact and reused as a component across my security "
     "platform.", italic=True)

doc.add_heading("1. From research prototype to a shipped, reusable model", level=1)
para(
    "Exhibit 11 introduced RegMap: a fine-tuned sentence-embedding model that maps NIST SP 800-53 "
    "security controls to the most relevant HIPAA Security Rule provisions. Since then I have taken "
    "RegMap from a paper artifact to a packaged, publicly usable model and made it a working "
    "component of a larger operational platform. This addendum records that progression, which "
    "speaks directly to being well positioned to advance the endeavor and to disseminating results "
    "for public benefit.")

doc.add_heading("2. Packaged for out-of-the-box public use", level=1)
para("RegMap is released in the standard Hugging Face / sentence-transformers format so any "
     "practitioner can use it immediately, with an honest model card and an open license:")
table(["Release element", "Detail"],
      [["Model", "Fine-tuned all-MiniLM-L6-v2 (384-dim), MultipleNegativesRankingLoss on curated "
        "NIST↔HIPAA pairs; loads with a single line of code."],
       ["Bundled corpus + wrapper", "The HIPAA provision corpus and a small inference helper "
        "(map_control → top-k HIPAA citations) ship with the model, so the NIST→HIPAA mapping works "
        "out of the box."],
       ["Model card", "States intended use, training data, evaluation, and limitations — RegMap is "
        "an assistive top-k retriever (correct provision in the top-5 ≈ 74% of the time) for an "
        "expert to confirm, not an authoritative single-answer classifier."],
       ["License", "Apache-2.0 (inherited from the base model), enabling free public and commercial "
        "use — reducing dependence on costly proprietary compliance tooling."],
       ["Distribution (published)", "Released publicly on the Hugging Face Hub "
        "(huggingface.co/stetteh/regmap-embedder), as a downloadable GitHub Release "
        "(v0.1-regmap), and as a runnable Docker image "
        "(ghcr.io/samuelgtetteh/regmap-embedder) that serves the mapping over an API."]])
doc.add_heading("2a. Public availability", level=1)
para("RegMap is now publicly available for anyone to use immediately, through three channels:")
bullet("Hugging Face Hub — https://huggingface.co/stetteh/regmap-embedder (load with one line of "
       "sentence-transformers).")
bullet("GitHub Release v0.1-regmap — a self-contained archive (model + HIPAA corpus + inference "
       "wrapper) that runs offline.")
bullet("Docker image — ghcr.io/samuelgtetteh/regmap-embedder:0.1, a serving container exposing a "
       "POST /map endpoint that returns the top-ranked HIPAA provisions for a NIST control.")

doc.add_heading("3. Reused as a component across the platform", level=1)
para("RegMap is not a one-off: the same model now powers multiple capabilities of the operational "
     "system, demonstrating durable, transferable value:")
bullet("Assessment tool — RegMap maps discovered assets/services to the NIST 800-53 controls an "
       "organization should implement (see Exhibit 18).")
bullet("Compliance advisor interview — the same embedder interprets a user's plain-English answers, "
       "matching them to the correct environment characteristics without brittle keyword rules.")
bullet("Real-time triage — the retrieval approach that RegMap validated is reused to surface the "
       "most relevant compliance controls for a live alert (Exhibit 16).")

doc.add_heading("4. Significance", level=1)
para("Releasing RegMap as an open, documented, reusable model — and embedding it as infrastructure "
     "across a working security platform — shows the proposed endeavor producing artifacts that "
     "others in the United States can adopt directly, and shows me carrying research through to "
     "dissemination and operational reuse. The full model, release tooling, and model card are "
     "public:")
p = doc.add_paragraph(); p.add_run(REPO).bold = True

para("")
para("Date: " + RUN_DATE)
para("Prepared by: Samuel Gbli Tetteh")

out = r"c:\Users\User\ai-cybersecurity-portfolio\exhibits\Exhibit 11A RegMap Reusable Published Model.docx"
doc.save(out)
print("saved:", out)
print("paragraphs:", len(doc.paragraphs), "| tables:", len(doc.tables))
