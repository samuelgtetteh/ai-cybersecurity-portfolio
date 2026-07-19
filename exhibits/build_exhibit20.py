"""Generate "Exhibit 20 - Testing Infrastructure: Target Labs for Continuous Validation.docx"."""
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

RUN_DATE = "July 17, 2026"
REPO = "https://github.com/samuelgtetteh/ai-cybersecurity-portfolio"
SHOTS = os.path.join(os.path.dirname(__file__), "evidence", "shots")   # tidy copies (see crop_shots.py)

doc = Document()

# --- Document-wide typography -------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
npf = normal.paragraph_format
npf.space_after = Pt(6)
npf.line_spacing = 1.15
npf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

h1 = doc.styles["Heading 1"]
h1.font.name = "Calibri"
h1.font.size = Pt(13)
h1.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
h1.paragraph_format.space_before = Pt(12)
h1.paragraph_format.space_after = Pt(4)
h1.paragraph_format.keep_with_next = True


def para(t, bold=False, italic=False):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = bold; r.italic = italic; return p


def bullet(t):
    doc.add_paragraph(t, style="List Bullet")


def figure(filename, caption):
    """Embed a screenshot (fit to page width) with a bold-labelled caption beneath it, kept together."""
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(8); pf.space_after = Pt(2); pf.keep_with_next = True
    path = os.path.join(SHOTS, filename)
    if os.path.exists(path):
        p.add_run().add_picture(path, width=Inches(6.3))
    else:
        p.add_run("[screenshot not found: %s]" % filename).italic = True
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cf = c.paragraph_format
    cf.space_after = Pt(12); cf.line_spacing = 1.0
    label, _, rest = caption.partition(" — ")
    lr = c.add_run(label + " — "); lr.bold = True; lr.font.size = Pt(9)
    rr = c.add_run(rest); rr.italic = True; rr.font.size = Pt(9)


def table(headers, rows):
    tb = doc.add_table(rows=1, cols=len(headers)); tb.style = "Light Grid Accent 1"
    tb.autofit = True
    for j, h in enumerate(headers):
        run = tb.rows[0].cells[j].paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(10)
    for row in rows:
        cells = tb.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = str(v)
            for pp in cells[j].paragraphs:
                pp.paragraph_format.space_after = Pt(2)
                pp.paragraph_format.line_spacing = 1.0
                for rn in pp.runs:
                    rn.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)   # breathing room after the table
    return tb


t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_after = Pt(2)
r = t.add_run("Exhibit 20 — Testing Infrastructure: Target Labs for Continuous Validation")
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(2)
sr = sub.add_run("Petitioner: Samuel Gbli Tetteh"); sr.bold = True
lede = doc.add_paragraph(); lede.alignment = WD_ALIGN_PARAGRAPH.CENTER
lede.paragraph_format.space_after = Pt(10)
lr = lede.add_run("NIW Proposed Endeavor — the purpose-built test harnesses I created to continuously "
                  "and reproducibly validate the detection and scanning systems, now themselves "
                  "released as open, standalone tools.")
lr.italic = True; lr.font.size = Pt(10.5)

doc.add_heading("1. Purpose", level=1)
para(
    "Building detectors is only half of credible security engineering; you must also prove they "
    "work, continuously and against realistic activity. To do that I built two disposable 'target "
    "labs' — controlled environments that generate realistic malicious and benign activity for the "
    "systems to detect. This exhibit documents them and their current phase: both are complete, in "
    "use, and now published as open tools so the validation is reproducible by anyone.")

doc.add_heading("2. The two labs", level=1)
table(["Lab", "What it does", "Validates"],
      [["Live Target Lab", "Two standing services that continuously stream synthetic-but-realistic "
        "identity login events and OT/ICS sensor readings to the live detection APIs (~15% "
        "intentionally malicious), and report the KNOWN injected label back through the feedback "
        "channel.", "The identity + OT/ICS detectors and the decision layer — end to end, live."],
       ["Cloud Target Lab", "Stands up a disposable fake-AWS (LocalStack) seeded with a deliberate "
        "mix of correctly-configured and intentionally-insecure resources (public S3 bucket, "
        "SSH-open security group, admin-* IAM user, etc.).", "The cloud security scanner / control "
        "auditing — with known ground-truth findings to detect."]])

doc.add_heading("3. Why this is rigorous validation", level=1)
bullet("Continuous, not one-shot: the live lab runs forever and generates fresh data each tick, "
       "closer to a real SIEM/SCADA stream than replaying a fixed dataset.")
bullet("Ground-truth feedback loop: because each generator knows the label it injected, it reports "
       "that label back, so the platform can compute LIVE precision/recall/specificity from its own "
       "trail (the same channel a real analyst or SOAR tool would use).")
bullet("Known-answer cloud fixtures: the cloud lab seeds a mix of secure and insecure resources, so "
       "a scanner is measured against genuine, varied findings — not an all-or-nothing fixture.")
bullet("Safe and disposable: everything runs in throwaway containers against emulated/synthetic "
       "targets — no real accounts, no third-party systems.")

doc.add_heading("4. Published as open, standalone tools", level=1)
para("Both labs are released so the validation is fully reproducible:")
table(["Lab", "GitHub", "Docker (GHCR)", "Hugging Face"],
      [["Live Target Lab", "github.com/samuelgtetteh/live-target-lab (Release v0.1)",
        "ghcr.io/samuelgtetteh/live-target-lab", "dataset stetteh/live-target-lab-events; "
        "Space stetteh/live-target-lab"],
       ["Cloud Target Lab", "github.com/samuelgtetteh/cloud-target-lab (Release v0.1)",
        "ghcr.io/samuelgtetteh/cloud-target-lab", "dataset stetteh/cloud-target-lab-scenarios; "
        "Space stetteh/cloud-target-lab"]])

doc.add_heading("5. Measured results from a continuous run", level=1)
para("The labs are not merely demonstrable — they have been run continuously against the live "
     "platform. In a representative run recorded July 2026, the platform and both event sources ran "
     "for approximately 22 hours without interruption (restart-persistent containers), streaming "
     "synthetic identity and OT/ICS events and reporting the injected ground-truth labels back. The "
     "following are measured live from the platform's own trail (GET /decision/stats and "
     "/decision/metrics) — not re-derived or estimated:")
table(["Quantity", "Value"],
      [["Continuous uptime", "~22 hours (backend + identity and OT/ICS event sources, all "
        "restart-persistent containers)"],
       ["Labelled verdicts retained", "100,005 of 100,009 — the trail reached its 100,000-record "
        "FIFO retention cap; ~100% labelled automatically via the feedback loop"],
       ["Volume by model", "identity: 66,201 labelled; OT/ICS: 33,824 labelled"],
       ["Overall precision / recall / specificity", "0.9961 / 0.9982 / 0.9965  (TP 47,083 · FP 183 · "
        "FN 84 · TN 52,666)"],
       ["Identity precision / recall / specificity", "1.000 / 0.998 / 1.000  (TP 42,015 · FP 0 · "
        "FN 84 · TN 24,102)"],
       ["OT/ICS precision / recall / specificity", "0.965 / 1.000 / 0.994  (TP 5,070 · FP 183 · "
        "FN 0 · TN 28,571)"],
       ["Response actions", "recorded by the Act layer (log / step-up-auth / webhook stubs) — visible "
        "in the console's Response Actions panel"]])
para("These figures are shown directly in Section 6 below, captured from the running console and its "
     "own API endpoints.")
para("This one run validates several things simultaneously: the detectors' live accuracy against a "
     "continuous, labelled stream at scale; the ground-truth feedback loop (which let these metrics "
     "be computed directly from the database rather than tallied by hand, as in Exhibit 14); the FIFO "
     "retention design (the trail held at its ~100,000-record cap over 22 hours instead of growing "
     "without bound); and restart-persistence. It is the testing infrastructure demonstrating the "
     "very properties it was built to verify.", italic=True)

doc.add_heading("6. Screenshot evidence (captured from the live system)", level=1)
para("The following screenshots were captured on July 17, 2026 directly from the running platform at "
     "the ~22-hour mark; they are the source of the numbers in Section 5. Figures 1–2 are the live "
     "Monitor console. Figures 3–6 are the platform's own JSON API responses (the ground truth behind "
     "the dashboard, each shown with its endpoint URL) so the results can be read raw and "
     "independently verified.")

figure("fig1_monitor_1031.png",
       "Figure 1 — Live Monitor console at 10:31:41 (100,042 events). Per-model panels: Identity "
       "66,239 processed, precision 100.0% / recall 99.8%; OT/ICS 33,801 processed, precision 96.5% / "
       "recall 100.0%. The live decision feed, alert queue, and recorded response actions are all "
       "visible — the full Record → Decide → Act → alert pipeline running end to end.")

figure("fig2_monitor_1032.png",
       "Figure 2 — The same console 49 seconds later at 10:32:30: the event counter has advanced "
       "(100,042 → 100,096; Identity 66,239 → 66,278; OT/ICS 33,801 → 33,816) and new decisions have "
       "streamed in — demonstrating the system is genuinely live and continuously processing, not a "
       "static snapshot.")

figure("fig3_stats.png",
       "Figure 3 — Raw GET /decision/stats: verdicts 100,009, labeled 100,005, by model "
       "(ics 33,813 / identity 66,194 / scan 2), retention cap max_verdicts 100,000. Confirms the "
       "FIFO retention design is holding the trail at its cap over the long run.")

figure("fig4_metrics_all.png",
       "Figure 4 — Raw GET /decision/metrics (all models): TP 47,083 · FP 183 · FN 84 · TN 52,666, "
       "precision 0.9961, recall 0.9982, specificity 0.9965, accuracy 0.9973 — computed by the "
       "platform from its own ground-truth-labelled trail.")

figure("fig5_metrics_identity.png",
       "Figure 5 — Raw GET /decision/metrics?model=identity: TP 42,015 · FP 0 · FN 84 · TN 24,102, "
       "precision 1.000, recall 0.998, specificity 1.000 — zero false positives over ~66,000 "
       "labelled identity events.")

figure("fig6_metrics_ics.png",
       "Figure 6 — Raw GET /decision/metrics?model=ics: TP 5,070 · FP 183 · FN 0 · TN 28,571, "
       "precision 0.965, recall 1.000, specificity 0.994 — zero missed attacks over ~33,800 labelled "
       "OT/ICS readings.")

doc.add_heading("7. Significance", level=1)
para("Purpose-built, continuously-running, ground-truth-labelled test harnesses — released openly "
     "alongside the systems they validate — demonstrate the engineering and scientific maturity to "
     "not only build security capabilities but to prove and reproduce their effectiveness. This "
     "supports both the credibility of the results in Exhibits 11–19 and that the petitioner is well "
     "positioned to advance the endeavor to a rigorous, verifiable standard.")
repo = doc.add_paragraph(); repo.alignment = WD_ALIGN_PARAGRAPH.LEFT
repo.paragraph_format.space_before = Pt(6)
repo.add_run("Repository: ").bold = True
repo.add_run(REPO)
foot = doc.add_paragraph(); foot.alignment = WD_ALIGN_PARAGRAPH.LEFT
foot.paragraph_format.space_before = Pt(10)
foot.add_run("Date: " + RUN_DATE + "\nPrepared by: Samuel Gbli Tetteh")

out = r"c:\Users\User\ai-cybersecurity-portfolio\exhibits\Exhibit 20 Testing Infrastructure Target Labs.docx"
doc.save(out); print("saved:", out, "| paras", len(doc.paragraphs), "| tables", len(doc.tables))
