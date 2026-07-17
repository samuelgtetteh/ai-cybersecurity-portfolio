"""
Generate "Exhibit 18 - Vulnerability-to-Compliance Assessment Tool.docx".

Documents the assessment capability built on top of the platform: a single workflow that discovers
assets (server scan, on-prem agent, or by ingesting another scanner's report), identifies and
prioritizes vulnerabilities (CVEs + CISA KEV + EPSS), and maps the findings to NIST SP 800-53
controls (using the RegMap model) with a conversational interview and downloadable reports. Mirrors
the Exhibit 14/16/17 style and describes the shipped, reproducible system.

Run:  venv\\Scripts\\python.exe exhibits\\build_exhibit18.py
"""
from docx import Document
from docx.shared import Pt

RUN_DATE = "July 16, 2026"
REPO = "https://github.com/samuelgtetteh/ai-cybersecurity-portfolio"

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = bold; r.italic = italic
    return p


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        run = t.rows[0].cells[j].paragraphs[0].add_run(h); run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = str(v)
    return t


title = doc.add_paragraph()
r = title.add_run("Exhibit 18: Vulnerability-to-Compliance Assessment Tool "
                  "(Discovery → CVE/KEV/EPSS → NIST SP 800-53)")
r.bold = True; r.font.size = Pt(15)
para("Petitioner: Samuel Gbli Tetteh", bold=True)
para("NIW Proposed Endeavor — Extending the platform across the full cyber-risk lifecycle: "
     "discovering assets, identifying and prioritizing their vulnerabilities, and mapping those "
     "findings to the exact security controls an organization must implement.", italic=True)

doc.add_heading("1. Purpose", level=1)
para(
    "Earlier exhibits covered detection (identity, OT/ICS), compliance mapping (RegMap), and the "
    "real-time decisioning platform. This exhibit documents the assessment tool I built to complete "
    "the lifecycle: a single browser workflow that goes from 'what do I have and what's wrong with "
    "it' (asset discovery + vulnerabilities) to 'what must I do about it' (prioritized NIST SP "
    "800-53 controls and a compliance report). It is designed to be usable by organizations that "
    "cannot afford commercial vulnerability-management suites.")

doc.add_heading("2. Design principle — orchestrate and ingest, then add the compliance layer", level=1)
para(
    "The tool does not attempt to replace mature vulnerability scanners such as Nessus or OpenVAS, "
    "which maintain large, curated detection feeds. Instead it orchestrates open-source discovery, "
    "ingests those scanners' reports, and adds the layer that is my own contribution: prioritization "
    "by real-world exploitability and mapping to NIST controls. This is both more achievable and more "
    "defensible than re-implementing vulnerability detection, and it interoperates with tools "
    "organizations already run.")

doc.add_heading("3. What it does", level=1)
table(["Stage", "Capability"],
      [["Discover", "Asset/port/service discovery via a built-in dependency-free scanner or nmap; "
        "for a private network that a hosted console cannot reach, a lightweight on-prem AGENT is "
        "run inside the network and reports results back; alternatively, ingest an existing Nessus / "
        "OpenVAS / nuclei report."],
       ["Identify", "Discovered services are mapped to known CVEs (NVD)."],
       ["Prioritize", "Each CVE is enriched with CISA KEV (is it being actively exploited?) and "
        "FIRST EPSS (probability of exploitation) and scored into a blended risk — surfacing what "
        "actually matters, not just raw CVSS."],
       ["Map to controls", "Findings are mapped to the relevant NIST SP 800-53 controls using the "
        "RegMap model (Exhibit 11/11A) — the vulnerability view feeds the compliance view."],
       ["Advise & report", "A conversational interview (local language model) captures the "
        "organization's context; the tool then produces prioritized control recommendations and "
        "downloadable DOCX / XLSX / JSON reports. The user may stop after discovery (scan-only) or "
        "continue to the full compliance package."]])

doc.add_heading("4. A single, unified workflow", level=1)
para("Discovery and compliance were consolidated into one tool with one scan engine: the user scans "
     "(or imports/agents), REVIEWS and verifies the findings, and then chooses whether to generate "
     "the compliance package. Everything runs from the browser — no command line, no separate "
     "products — which is the difference between a research demo and something a small security team "
     "or a state/local government office could actually operate.")

doc.add_heading("5. Safety and authorization", level=1)
para("Active scanning is gated: targets are restricted to authorized ranges by default, an explicit "
     "acknowledgement is required before each scan, and posture-changing response actions are "
     "recorded stubs rather than live changes. The AI is advisory only; a human confirms findings "
     "and decisions. These choices reflect responsible, authorized-use security engineering.")

doc.add_heading("6. Reproducibility", level=1)
para("The tool is public and reproducible:")
bullet("backend/securescan/ — discovery engines, the shared analyzer (CVE mapping + KEV/EPSS + "
       "service→control-category derivation), the engine catalog, and report importers.")
bullet("backend/agent/agent.py — the self-contained on-prem scan agent.")
bullet("backend/advisor_api.py — the unified scan → analyze → interview → report pipeline, reusing "
       "the RegMap model for control mapping. RegMap is itself publicly released — Hugging Face "
       "(huggingface.co/stetteh/regmap-embedder) and Docker (ghcr.io/samuelgtetteh/regmap-embedder) "
       "— so the control-mapping component is open and reproducible (see Exhibit 11A).")
bullet("tests/ — automated tests for discovery, enrichment, importers, the agent job lifecycle, and "
       "the interview, served from the same reproducible Docker image as the detectors.")
p = doc.add_paragraph(); p.add_run(REPO).bold = True

doc.add_heading("7. Significance for the “Well Positioned” Prong", level=1)
para("This exhibit shows the endeavor carried across the entire cyber-risk lifecycle — discovery, "
     "vulnerability identification, exploit-aware prioritization, and control prescription — in one "
     "accessible, open tool that reuses my own published research (the RegMap model) and interoperates "
     "with the scanners organizations already use. It demonstrates end-to-end capability across "
     "research, machine learning, systems engineering, and operational security, and it directly "
     "benefits U.S. organizations that lack the budget for commercial platforms. Together with "
     "Exhibits 11–17, it is direct evidence that I am well positioned to advance this endeavor.")

para("")
para("Date: " + RUN_DATE)
para("Prepared by: Samuel Gbli Tetteh")

out = r"c:\Users\User\ai-cybersecurity-portfolio\exhibits\Exhibit 18 Vulnerability-to-Compliance Assessment Tool.docx"
doc.save(out)
print("saved:", out)
print("paragraphs:", len(doc.paragraphs), "| tables:", len(doc.tables))
