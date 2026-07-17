"""
Generate "Exhibit 19 - Integrated Security Platform: Real-World Applications and National Impact.docx".

Ties the individual components (RegMap compliance mapping, hybrid-identity anomaly detection, OT/ICS
intrusion detection, the real-time decisioning platform, and the vulnerability-to-compliance
assessment tool) into one integrated platform, illustrated by a concrete healthcare attack-lifecycle
scenario, and maps each component to U.S. government cybersecurity priorities. Mirrors the
Exhibit 14/16/17/18 style.

Run:  venv\\Scripts\\python.exe exhibits\\build_exhibit19.py
"""
from docx import Document
from docx.shared import Pt

RUN_DATE = "July 16, 2026"
REPO = "https://github.com/samuelgtetteh/ai-cybersecurity-portfolio"

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = bold; r.italic = italic
    return p


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def numbered(text):
    doc.add_paragraph(text, style="List Number")


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        run = t.rows[0].cells[j].paragraphs[0].add_run(h); run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = str(v)
    return t


title = doc.add_paragraph()
r = title.add_run("Exhibit 19: Integrated Security Platform — Real-World Applications and National Impact")
r.bold = True; r.font.size = Pt(15)
para("Petitioner: Samuel Gbli Tetteh", bold=True)
para("NIW Proposed Endeavor — Demonstrating that the individual research prototypes form a single, "
     "integrated platform spanning the full security lifecycle, and mapping that platform to U.S. "
     "national cybersecurity priorities.", italic=True)

doc.add_heading("1. From separate prototypes to one platform", level=1)
para(
    "The preceding exhibits document distinct contributions: RegMap (NIST→HIPAA compliance mapping), "
    "hybrid-identity anomaly detection, OT/ICS intrusion detection, a real-time decisioning platform, "
    "an analyst case-management workflow, and a vulnerability-to-compliance assessment tool. This "
    "exhibit shows they are not siloed: together they cover the full lifecycle of security operations "
    "— predict, discover, detect, prioritize, respond, and prove compliance — in one system a single "
    "organization can run.")

doc.add_heading("2. The components and what each does in the real world", level=1)
table(["Component", "Real-world role", "Who uses it"],
      [["RegMap — compliance mapping", "Cross-walks a control set to HIPAA (and, by extension, other "
        "frameworks) so audits and gap analyses take minutes, not weeks.", "Compliance officers, CISOs"],
       ["Hybrid-identity anomaly detection", "Flags credential compromise, lateral movement, and "
        "insider threats across cloud + on-prem identity in real time.", "SOC analysts, IAM teams"],
       ["OT/ICS intrusion detection", "Detects cyber-physical attacks on industrial/building systems "
        "from sensor data without needing labeled attacks.", "Plant/OT operators, hospital facilities"],
       ["Assessment tool", "Discovers assets, identifies and prioritizes vulnerabilities (CVE + KEV + "
        "EPSS), and prescribes NIST 800-53 controls.", "IT/security teams, MSPs"],
       ["Decisioning platform", "Records, prioritizes, and acts on all of the above, with auditable "
        "responses and AI-assisted triage.", "SOC leads, incident responders"]])

doc.add_heading("3. How they work together — a healthcare attack lifecycle", level=1)
para("Consider a mid-sized U.S. healthcare provider (hybrid identity, cloud-hosted patient data, "
     "building-management OT, HIPAA obligations). The platform operates end to end:")
numbered("Discover — the assessment tool scans the network and finds an unpatched server exposing a "
         "vulnerable remote-access service.")
numbered("Identify & prioritize — it maps the service to a critical CVE and flags it as actively "
         "exploited (CISA KEV) with a high EPSS score, raising it to the top of the queue.")
numbered("Prescribe — using RegMap it recommends the specific NIST 800-53 controls (e.g. remote "
         "access, flaw remediation) and generates the compliance evidence.")
numbered("Detect — simultaneously, the identity detector flags a privileged account reaching that "
         "server from an unusual source at an odd hour (possible compromise + lateral movement).")
numbered("Detect (cyber-physical) — the OT/ICS model flags an anomaly in the building-management "
         "system, a possible diversion or second stage.")
numbered("Decide & act — the decisioning platform correlates these into one prioritized, audited "
         "alert with recommended controls and response actions, ready for the analyst to resolve.")
para("The same platform thus predicted (documented controls), discovered and prioritized "
     "(vulnerabilities), detected (identity + OT), and prescribed remediation (controls) — with an "
     "auditable trail throughout.", italic=True)

doc.add_heading("4. Alignment with U.S. national priorities", level=1)
table(["National priority", "How the platform contributes"],
      [["CISA Cross-Sector Cybersecurity Performance Goals (CPGs)", "Asset inventory, vulnerability "
        "management with exploit-aware prioritization, and account/anomaly monitoring."],
       ["NIST Cybersecurity Framework", "Covers Identify (assets/controls), Protect (control "
        "prescription), Detect (identity + OT anomaly detection), and Respond (decisioning + actions)."],
       ["DHS critical-infrastructure sectors", "Directly applicable to Healthcare & Public Health, "
        "Energy, Water, and Critical Manufacturing (OT/ICS)."],
       ["Accessibility / cost", "Open-source and reuses published models, so state/local governments "
        "and small organizations that cannot afford commercial suites can adopt it."]])

doc.add_heading("5. Significance for the National Interest", level=1)
para("This exhibit shows the proposed endeavor is not theoretical: I have built a functioning, "
     "integrated platform that addresses urgent U.S. cybersecurity needs across IT and OT domains, "
     "protects critical infrastructure, aligns with federal priorities, and — being open and "
     "low-cost — broadens access to capabilities usually reserved for well-resourced enterprises. "
     "Together with Exhibits 11–18, it is direct evidence of both the national importance of the "
     "endeavor and my being well positioned to advance it.")
p = doc.add_paragraph(); p.add_run(REPO).bold = True

para("")
para("Date: " + RUN_DATE)
para("Prepared by: Samuel Gbli Tetteh")

out = r"c:\Users\User\ai-cybersecurity-portfolio\exhibits\Exhibit 19 Integrated Security Platform.docx"
doc.save(out)
print("saved:", out)
print("paragraphs:", len(doc.paragraphs), "| tables:", len(doc.tables))
