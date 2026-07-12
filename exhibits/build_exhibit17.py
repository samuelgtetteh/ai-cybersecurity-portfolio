"""
Generate "Exhibit 17 - Analyst Case-Management Workflow.docx".

Documents the human-in-the-loop case-management layer built on top of the Record -> Decide ->
Act -> AI-triage platform of Exhibit 16. The alert queue is no longer read-only: from the live
browser console an analyst can open a case, inspect the evidence (the contributing verdicts),
review the subject's outcome history and AI triage, and take auditable decisions — acknowledge,
assign, resolve (true/false positive), suppress/allowlist a subject, add notes, or fire a manual
response action. Crucially, a resolution feeds ground truth back to the contributing verdicts,
which trains the Decide layer's outcome-weighting: the human's decision measurably improves
future decisions. Mirrors the Exhibit 14/16 style and reports what the shipped system does.

Run:  venv\\Scripts\\python.exe exhibits\\build_exhibit17.py
"""
from docx import Document
from docx.shared import Pt

RUN_DATE = "July 12, 2026"
REPO = "https://github.com/samuelgtetteh/ai-cybersecurity-portfolio"

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = bold; r.italic = italic
    return p


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        run = t.rows[0].cells[j].paragraphs[0].add_run(h); run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = str(v)
    return t


# ---- title block ----
title = doc.add_paragraph()
r = title.add_run("Exhibit 17: Analyst Case-Management Workflow "
                  "(Human-in-the-Loop Decisioning with a Learning Feedback Loop)")
r.bold = True; r.font.size = Pt(15)
para("Petitioner: Samuel Gbli Tetteh", bold=True)
para("NIW Proposed Endeavor — Elevating the real-time decisioning platform (Exhibit 16) into an "
     "operational analyst workstation: turning raised alerts into managed cases that a human "
     "investigates and resolves, with each resolution feeding back to improve the system.",
     italic=True)

# 1. Purpose
doc.add_heading("1. Purpose", level=1)
para(
    "Exhibit 16 documented the platform that records, decides, acts, and triages. A detection "
    "platform, however, is only operational if a human analyst can act on its output. This exhibit "
    "documents the case-management workflow I built directly into the live browser console so that "
    "an analyst can open an alert, see the underlying evidence, and make and record a decision — "
    "without touching the API, the command line, or the database. It closes the last gap between "
    "an automated detector and a usable security-operations tool, and it does so with a design "
    "principle throughout: the human's decision is authoritative and is fed back to improve the "
    "system, while the language model remains strictly advisory.")

# 2. From a read-only queue to a managed case
doc.add_heading("2. From a Read-Only Queue to a Managed Case", level=1)
para("Previously the alert queue supported only two operations — view the AI rationale and close. "
     "The workflow now supports the full lifecycle of a security case. Clicking an alert opens a "
     "case drawer that presents the evidence and offers a decision toolbar:")
table(["Capability", "What the analyst can do", "Endpoint"],
      [["Open / drill-down",
        "Inspect the contributing verdicts (the actual events), the subject's confirmed "
        "malicious/benign history, the responder actions already taken, and the full case audit "
        "trail.",
        "GET /decision/alerts/{id}"],
       ["Acknowledge / assign",
        "Take ownership of a case (in-progress) or assign it to a named analyst; the case stays "
        "visible in the active queue.",
        "POST /decision/alerts/{id}/acknowledge, /assign"],
       ["Resolve (disposition)",
        "Close the case as true positive, false positive, or benign — which writes ground truth "
        "back to every contributing verdict (see section 3).",
        "POST /decision/alerts/{id}/resolve"],
       ["Suppress / allowlist",
        "Mute a confirmed known-good subject (e.g. a service account) for a chosen window; the "
        "Decide layer consults the allowlist before raising that subject again.",
        "POST /decision/alerts/{id}/suppress; GET/DELETE /decision/suppressions"],
       ["Notes / audit trail",
        "Attach free-text notes; every action (who, what, when, why) is journalled to an "
        "immutable case history.",
        "POST /decision/alerts/{id}/note"],
       ["Manual response",
        "Fire a responder on demand — ticket, webhook, or a posture-change action "
        "(disable account, require step-up auth).",
        "POST /decision/alerts/{id}/act"]])

# 3. The learning feedback loop
doc.add_heading("3. The Decision Trains the System", level=1)
para("The defining feature of the workflow is that an analyst's resolution is not merely filed — "
     "it improves future decisions. When a case is resolved as a true positive, every verdict "
     "that contributed to the alert is labelled malicious; when it is resolved as a false "
     "positive, they are labelled benign. That ground truth flows into the same outcome-weighting "
     "the Decide layer already uses (Exhibit 16, section 4): a subject repeatedly confirmed benign "
     "is thereafter suppressed as a chronic false positive, while one confirmed malicious is "
     "escalated. The human closes the loop, and the alert queue becomes more trustworthy the more "
     "it is used. This reuses the platform's existing feedback channel — the same mechanism the "
     "live event sources use to report ground truth — now driven from the analyst's console.")
para("This is also why the workflow keeps the language model advisory only. The model assists "
     "triage (it summarizes and prioritizes), but it never resolves a case or executes a "
     "posture-changing action. Those decisions are reserved for the human, both because the model "
     "reads attacker-controllable fields (a prompt-injection surface) and because accountable "
     "security decisions must rest with an accountable operator.", italic=True)

# 4. Safe-by-design response actions
doc.add_heading("4. Safe-by-Design Response Actions", level=1)
para("The manual response menu includes posture-changing actions — disable account and require "
     "step-up authentication — that a premium product must offer. In this deployment they are "
     "implemented as recorded stubs: the intent and target are written to the audit trail, but no "
     "live directory/identity-provider integration is wired, so the console cannot be induced "
     "(by a malicious input or a model error) to lock out a real account. The seam where a "
     "production deployment would call Azure AD / Okta / ServiceNow is explicit in the code. The "
     "console is therefore fully featured without becoming an attack surface.")

# 5. Everything from the browser
doc.add_heading("5. Operable Entirely from the Browser", level=1)
para("The whole workflow is driven from the live Server-Sent-Events console that replaced the raw "
     "API view: the queue updates in real time, and each case is opened and resolved in a drawer "
     "with the evidence in front of the analyst. No command-line access, no direct database "
     "queries, and no API client are required to run a security investigation end to end. This is "
     "the first stage of a broader effort to make the entire platform configurable and operable "
     "from the browser (planned: an in-console settings panel for the system's operational "
     "limits, and browser-triggered compliance scans).")

# 6. Reproducibility
doc.add_heading("6. Reproducibility and Verifiability", level=1)
para("The workflow is public and independently reproducible:")
bullet("backend/decision_api.py — the analyst endpoints (alert detail, acknowledge, assign, "
       "resolve, note, suppress, act, suppressions management).")
bullet("backend/verdict_store.py — the alert lifecycle fields, the alert_events audit trail, the "
       "suppression/allowlist store, and label_alert_verdicts (the loop-closing feedback).")
bullet("backend/policy.py — honours the allowlist before raising a subject's alert.")
bullet("backend/actions.py — the manual responders, including the safe posture-change stubs.")
bullet("backend/dashboard/index.html — the case drawer, decision toolbar, evidence view, and "
       "audit trail in the live console.")
bullet("tests/test_decision_layer.py — automated tests covering alert drill-down, the "
       "acknowledge/assign/note journal, true-positive resolution labelling the evidence, "
       "suppression blocking a re-alert, and the recorded manual-action stub.")
p = doc.add_paragraph(); p.add_run(REPO).bold = True

# 7. Significance
doc.add_heading("7. Significance for the “Well Positioned” Prong", level=1)
para("This exhibit shows that I have taken the endeavor beyond detection and automated "
     "decisioning into the human-facing operations layer that determines whether such a system is "
     "actually usable in a security operations center: a case-management workflow with evidence, "
     "auditable decisions, safe response actions, and — distinctively — a feedback loop in which "
     "the analyst's judgment continuously improves the model's future decisions. It demonstrates "
     "product-level systems and UX engineering on top of my own research artifacts, and it "
     "reflects mature security judgment (advisory-only AI, safe-by-design actions, full "
     "auditability). Together with Exhibits 11–16, it is direct evidence that I am well "
     "positioned to advance this endeavor independently.")

para("")
para("Date: " + RUN_DATE)
para("Prepared by: Samuel Gbli Tetteh")

out = r"c:\Users\User\ai-cybersecurity-portfolio\exhibits\Exhibit 17 Analyst Case-Management Workflow.docx"
doc.save(out)
print("saved:", out)
print("paragraphs:", len(doc.paragraphs), "| tables:", len(doc.tables))
