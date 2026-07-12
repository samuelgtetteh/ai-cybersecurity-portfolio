"""
Apply "under peer review" update notes to the existing NIW exhibits (option a).

- Exhibits 11, 12, 13: inserts a dated addendum immediately after the Section 1
  (Project Summary) body, recording that the prototype has been developed into a
  manuscript now under peer review, and cross-referencing Exhibit 15.
- "Project Include in document.docx" (third-person petition narrative): inserts a
  milestone paragraph at the end of the "Significance of the Three Projects" section.

Safe + idempotent:
  * copies each original to exhibits/_pre_review_update_backup/ before modifying;
  * skips a file that already contains an "Update (July 2026)" note, so re-running
    does not duplicate.

Run:  venv\\Scripts\\python.exe exhibits\\apply_review_updates.py
"""
import shutil
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

EX = Path(r"c:\Users\User\ai-cybersecurity-portfolio\exhibits")
BACKUP = EX / "_pre_review_update_backup"
BACKUP.mkdir(exist_ok=True)
MARK = "Update (July 2026)"

# per-exhibit addendum text (first person, matching Exhibits 11-14 voice)
NOTE = {
    "Exhibit11_RegMap_Project.docx": (
        "This work has since been developed into a full scientific manuscript, "
        "“Automated Regulatory Crosswalking: Fine-Tuned Semantic Retrieval for "
        "NIST SP 800-53 to HIPAA Mapping,” which is now under peer review at ACM "
        "Digital Threats: Research and Practice (DTRAP). The manuscript adds a rigorous "
        "retrieval evaluation—Recall@k, MRR, MAP, and nDCG with bootstrap confidence "
        "intervals and a stricter, leakage-free group-split protocol—and positions the "
        "work against the current compliance-mapping literature. A consolidated record of "
        "this and the related research progress is provided in Exhibit 15."),
    "Exhibit 12_hybrid _indentity.md.docx": (
        "This work has since been developed into a full scientific manuscript, "
        "“Access Breadth as a Robust Signal for Credential-Based Lateral Movement: A "
        "Red-Team Feature-Attribution Study on the LANL Dataset,” which is now under peer "
        "review at Transactions on Machine Learning Research (TMLR). The manuscript evaluates "
        "the detector against the LANL red-team ground truth—a rare source of real, "
        "labelled attack activity—with bootstrap confidence intervals and a full "
        "feature-attribution analysis. A consolidated record of this and the related research "
        "progress is provided in Exhibit 15."),
    "Exhibit 13 OT ICS Intrusion Detection.docx": (
        "This work has since been developed into a full scientific manuscript, "
        "“Label Leakage in ICS Anomaly Detection: A Reproducible Re-Evaluation of an "
        "Autoencoder Detector on the HAI Testbed,” which is now under peer review at ACM "
        "Digital Threats: Research and Practice (DTRAP). The manuscript formalizes the "
        "leakage audit and strict session-disjoint re-evaluation summarized in Section 4.4. A "
        "consolidated record of this and the related research progress is provided in "
        "Exhibit 15."),
}

NARRATIVE_FILE = "Project Include in document.docx"
NARRATIVE_NOTE = (
    "Each of these three prototypes has since been developed into an original scientific "
    "manuscript, and all three are now under peer review at established venues—the "
    "identity anomaly-detection study at Transactions on Machine Learning Research (TMLR), "
    "and the compliance-mapping and OT/ICS studies at ACM Digital Threats: Research and "
    "Practice (DTRAP). This advancement—from prototype, to live-validated system, to "
    "manuscript under external scientific review—further establishes that Mr. Tetteh is "
    "not merely planning the proposed endeavor but is actively and successfully executing it. "
    "A consolidated record of this progress and of the full body of work is provided in "
    "Exhibit 15.")


def has_mark(doc):
    return any(MARK in p.text for p in doc.paragraphs)


def insert_after(anchor: Paragraph, lead: str, body: str) -> Paragraph:
    """Insert a new italic paragraph right after `anchor`, with a bold lead run."""
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    np = Paragraph(new_p, anchor._parent)
    r1 = np.add_run(lead)
    r1.bold = True
    r1.italic = True
    r2 = np.add_run(body)
    r2.italic = True
    return np


def find_summary_body(doc):
    """Return the paragraph that is the body of Section 1 (Project Summary)."""
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        if p.text.strip() == "1. Project Summary":
            for j in range(i + 1, len(paras)):
                if paras[j].text.strip():
                    return paras[j]
    return None


def find_startswith(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    return None


def process(fname, anchor_fn, lead, body):
    src = EX / fname
    doc = Document(str(src))
    if has_mark(doc):
        print(f"SKIP (already updated): {fname}")
        return
    anchor = anchor_fn(doc)
    if anchor is None:
        print(f"ANCHOR NOT FOUND: {fname}  <-- not modified")
        return
    shutil.copy2(src, BACKUP / fname)
    insert_after(anchor, lead, body)
    doc.save(str(src))
    print(f"UPDATED: {fname}")


if __name__ == "__main__":
    for fname, body in NOTE.items():
        process(fname, find_summary_body, f"{MARK}: ", body)
    # narrative: anchor on the "Taken together" significance paragraph
    process(NARRATIVE_FILE,
            lambda d: find_startswith(d, "Taken together"),
            f"{MARK}: ", NARRATIVE_NOTE)
    print(f"\nBackups of originals saved in: {BACKUP}")
