"""
Generate "Exhibit 15 - Research Progress and Body of Work.docx".

A cumulative NIW evidence exhibit that (a) quantifies the full body of work built
across the three proposed-endeavor pathways and (b) documents its advancement from
live-validated prototypes (Exhibits 11-14) to three scientific manuscripts now under
peer review. Every figure is verifiable from the public GitHub repository as of the
"as of" date below; nothing here is projected or aspirational.

All quantitative metrics were measured directly from the repository on 2026-07-11:
  - 31 commits over 9 active days, 2026-06-08 to 2026-07-11
  - 78 version-controlled files
  - 48 Python modules, 8,037 lines of Python (excluding third-party dependencies)
  - 10 Jupyter notebooks, ~2,052 lines of executable code
  - 3 manuscripts: ~7,675 + ~4,837 + ~3,516 = ~16,000 words; 32 + 27 + 29 = 88 reference entries

Re-run:  venv\\Scripts\\python.exe exhibits\\build_exhibit15.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

AS_OF = "July 11, 2026"
REPO = "https://github.com/samuelgtetteh/ai-cybersecurity-portfolio"

doc = Document()

# base font
st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(11)


def h(text, level):
    p = doc.add_heading(text, level=level)
    return p


def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for j, hd in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        run = c.paragraphs[0].add_run(hd)
        run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = str(val)
    return t

# ---------------------------------------------------------------- title block
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = title.add_run("Exhibit 15: Summary of Research Progress and Body of Work")
r.bold = True
r.font.size = Pt(15)

para("Petitioner: Samuel Gbli Tetteh", bold=True)
para("NIW Proposed Endeavor — Cumulative evidence across all three research "
     "pathways: a documented record of sustained, independent execution and of the "
     "advancement of that work from validated prototypes to peer-reviewed scientific "
     "research.", italic=True)

# ---------------------------------------------------------------- 1. Purpose
h("1. Purpose", 1)
para(
    "Exhibits 11–14 documented three functional research prototypes—one for "
    "each pathway of my proposed endeavor—and their validation as live, "
    "containerized services. This exhibit serves two further purposes. First, it "
    "consolidates the full body of work into a single, quantified record so that the "
    "scope and sustained nature of my independent effort can be assessed at a glance. "
    "Second, and most importantly, it documents a decisive advancement that post-dates "
    "those earlier exhibits: each of the three prototype systems has since been "
    "developed into a rigorous scientific manuscript, and all three are now under peer "
    "review at established venues. This progression—from working prototype, to "
    "live-validated system, to peer-reviewed scholarship—is direct evidence for the "
    "second prong of Matter of Dhanasar: that I am well positioned to advance the "
    "proposed endeavor. It reflects not an intention to conduct research, but a "
    "demonstrated and continuing record of executing it.")

# ------------------------------------------- 2. headline: papers under review
h("2. Principal Advancement: Three Manuscripts Under Peer Review", 1)
para(
    "In July 2026 I completed and submitted three original scientific manuscripts for "
    "peer review, one corresponding to each research pathway. Each manuscript takes the "
    "prototype documented in the referenced exhibit and subjects it to a rigorous, "
    "independently reproducible scientific evaluation. Their submission moves this work "
    "from the category of engineering prototypes into the category of contributions to "
    "the peer-reviewed scholarly record.")
table(
    ["Manuscript", "Pathway", "Venue (peer-reviewed)", "Status (as of " + AS_OF + ")"],
    [
        ["Access Breadth as a Robust Signal for Credential-Based Lateral Movement: "
         "A Red-Team Feature-Attribution Study on the LANL Dataset",
         "Pathway 1 (Exhibit 12)",
         "Transactions on Machine Learning Research (TMLR)", "Submitted; under review"],
        ["Label Leakage in ICS Anomaly Detection: A Reproducible Re-Evaluation of an "
         "Autoencoder Detector on the HAI Testbed",
         "Pathway 2 (Exhibit 13)",
         "ACM Digital Threats: Research and Practice (DTRAP)", "Submitted; under review"],
        ["Automated Regulatory Crosswalking: Fine-Tuned Semantic Retrieval for "
         "NIST SP 800-53 to HIPAA Mapping",
         "Pathway 3 (Exhibit 11)",
         "ACM Digital Threats: Research and Practice (DTRAP)", "Submitted; under review"],
    ])
para(
    "Transactions on Machine Learning Research is a peer-reviewed journal for machine "
    "learning research; ACM Digital Threats: Research and Practice is a peer-reviewed "
    "journal of the Association for Computing Machinery dedicated to the science and "
    "practice of countering digital threats. Submitting to venues of this standing, and "
    "preparing each manuscript to their double-anonymous review requirements, is itself "
    "evidence of research maturity and of engagement with the professional scientific "
    "community central to the national interest my endeavor serves.")

# ------------------------------------------------- 3. quantitative body of work
h("3. Quantitative Summary of the Body of Work", 1)
para(
    "The following figures were measured directly from the public repository as of "
    + AS_OF + ". They are fully verifiable: the repository, its complete version-control "
    "history, and every artifact below are publicly accessible.")
table(
    ["Dimension", "Measured extent"],
    [
        ["Active development period", "2026-06-08 to 2026-07-11 (building on prior academic work)"],
        ["Version-control commits", "31 commits across 9 distinct active days"],
        ["Version-controlled files", "78 files"],
        ["Python source", "48 modules; over 8,000 lines (excluding third-party dependencies)"],
        ["Jupyter notebooks", "10 documented, executable notebooks (~2,050 lines of code)"],
        ["Scientific manuscripts", "3 manuscripts; ~16,000 words; 88 cited-reference entries"],
        ["Trained / fine-tuned models",
         "Deep autoencoder (OT/ICS); Isolation Forest (identity); fine-tuned "
         "Sentence-BERT retriever (compliance); plus preserved leaked/strict autoencoder "
         "variants retained for the reproducibility audit"],
        ["Public benchmark datasets used",
         "HAI ICS testbed; LANL authentication + red-team logs; NIST SP 800-53 – HIPAA crosswalk"],
        ["Deployed systems",
         "One containerized multi-model REST API serving all three detectors; one "
         "natural-language Control Advisor CLI; two supporting live-simulation repositories"],
        ["Supporting documentation",
         "Four prior evidence exhibits (11–14); reproducibility runbooks; dated progress log"],
    ])

# ---------------------------------------------------------------- 4. timeline
h("4. Development Timeline", 1)
para(
    "The chronology below is drawn from the repository's version-control history and "
    "documents continuous, self-directed progress over the period.")
table(
    ["Date (2026)", "Milestone"],
    [
        ["Jun 8", "Project structure initialized"],
        ["Jun 15–17", "Pathway 3 (compliance): NIST–HIPAA data preparation; "
         "Sentence-BERT model trained and evaluated; demonstration interface"],
        ["Jun 25", "Compliance crosswalk data pipeline finalized"],
        ["Jul 3", "Pathway 2 (OT/ICS): deep autoencoder built on the HAI testbed; "
         "Pathway 1 (identity) notebook on 2M LANL events; Exhibit 13 prepared"],
        ["Jul 4", "All three models engineered into live REST APIs; OT/ICS data-leak "
         "identified and corrected; Control Advisor tool built"],
        ["Jul 4–5", "System landscape documented; live-testing environment, tally "
         "tooling, automated test suite, and backend hardening"],
        ["Jul 5", "Live-environment validation captured (Exhibit 14); scientific-paper "
         "program planned; first drafts of all three papers"],
        ["Jul 10–11", "Rigorous extended re-evaluations: exact label-leakage "
         "quantification, session-disjoint retraining, determinism fixes, group-split and "
         "nDCG robustness; three self-contained reproducible evaluation notebooks; all "
         "three manuscripts expanded to scholarly standard"],
        ["Jul 11", "All three manuscripts submitted for peer review (Pathway 1 → TMLR; "
         "Pathways 2 and 3 → ACM DTRAP)"],
    ])

# ------------------------------------------- 5. per-pathway advancement
h("5. Advancement of Each Pathway: Prototype → Rigorous Science → Peer Review", 1)

h("5.1 Pathway 1 — Context-Aware Identity Anomaly Detection", 2)
para(
    "The prototype in Exhibit 12 detected anomalies across two million LANL "
    "authentication events. It has since been advanced into a controlled scientific "
    "study evaluated against the LANL red-team ground truth—a rare source of real, "
    "labelled attack activity—with bootstrap confidence intervals and a full "
    "feature-attribution analysis identifying access breadth as the robust, interpretable "
    "signal that drives detection. That study is the manuscript now under review at TMLR.")

h("5.2 Pathway 2 — OT/ICS Intrusion Detection", 2)
para(
    "The Exhibit 13 autoencoder was subjected to a full audit of my own pipeline. The "
    "audit exactly quantified a label-leakage flaw (a literally perfect, and therefore "
    "meaningless, ROC AUC of 1.000000), demonstrated the contamination to be model-deep, "
    "identified a second cross-session contamination flaw, and re-evaluated the detector "
    "under a strict session-disjoint protocol with calibration and alarm-aggregation "
    "analyses. The resulting reproducibility study is the manuscript now under review at "
    "ACM DTRAP.")

h("5.3 Pathway 3 — Automated Compliance Validation", 2)
para(
    "The RegMap prototype in Exhibit 11 was developed into a rigorous retrieval study of "
    "NIST SP 800-53 → HIPAA mapping, reporting Recall@k, MRR, MAP, and nDCG with "
    "bootstrap confidence intervals, a stricter leakage-free group-split protocol, and an "
    "explicit positioning against the current compliance-mapping literature. That study is "
    "the manuscript now under review at ACM DTRAP.")

# ------------------------------------------- 6. rigor / maturity
h("6. Demonstrated Scientific Rigor and Research Maturity", 1)
para(
    "A distinguishing feature of this body of work is that its rigor is self-imposed. In "
    "preparing the OT/ICS manuscript I audited my own earlier, more favorable results, "
    "discovered that they were inflated by a data leak, preserved rather than discarded "
    "the flawed artifacts, and reported the corrected numbers as an explicit ‘honesty "
    "ladder’ (leaked 1.000 → pooled 0.929 → session-disjoint 0.869). In the "
    "identity work I identified and fixed two reproducibility defects that had been "
    "silently varying results, and re-derived every reported number deterministically. "
    "Finding, disclosing, and exactly quantifying the limitations of one’s own work—"
    "rather than reporting only the most flattering result—is the standard of scientific "
    "integrity expected of an independent researcher, and it is evidenced throughout this "
    "portfolio.")

# ------------------------------------------- 7. reproducibility
h("7. Reproducibility and Verifiability", 1)
para(
    "Every claim in this exhibit is independently verifiable. All source code, notebooks, "
    "trained models, evaluation scripts, and documentation are publicly available at:")
p = doc.add_paragraph()
p.add_run(REPO).bold = True
para(
    "Each reported result is reproducible from tracked, self-contained evaluation "
    "notebooks; all datasets are public; and the complete version-control history "
    "substantiates the timeline in Section 4. No proprietary tools, data, or "
    "employer resources were used at any stage.")

# ------------------------------------------- 8. significance
h("8. Significance for the ‘Well Positioned’ Prong", 1)
para(
    "Taken together, this record establishes that I am well positioned to advance the "
    "proposed endeavor without an employer-sponsored position:")
bullet("I have independently executed the full three-pathway research program I "
       "proposed—not as concepts, but as built, measured, and publicly released systems.")
bullet("I have sustained that effort over a continuous period, producing a substantial "
       "and verifiable body of software, models, and documentation.")
bullet("I have demonstrated the engineering maturity to unify independent prototypes into "
       "a single deployed, live-validated system (Exhibit 14).")
bullet("I have demonstrated the scientific maturity to audit, correct, and rigorously "
       "re-evaluate my own work to publication standard.")
bullet("Most significantly, I have advanced all three efforts into original manuscripts "
       "now under peer review at established venues, submitting my work for external "
       "scientific scrutiny—the clearest available evidence of a genuine and continuing "
       "record of progress.")
h("Continued Progress Since Submission: An Operational Platform", 1)
para(
    "Since submitting the three manuscripts, I have continued the endeavor by building the "
    "prototypes into a single operational security platform, all publicly reproducible:")
bullet("Real-time decisioning platform — the detectors now Record every decision, Decide via a "
       "policy engine, Act with auditable responses, and produce AI-assisted triage (Exhibit 16).")
bullet("Analyst case-management workflow — a human-in-the-loop console where resolving an alert "
       "feeds ground truth back to improve future decisions (Exhibit 17).")
bullet("Vulnerability-to-compliance assessment tool — one workflow from asset discovery and "
       "exploit-aware vulnerability prioritization (CVE + CISA KEV + EPSS) to NIST SP 800-53 control "
       "recommendations and reports, including an on-prem scan agent and ingestion of Nessus/OpenVAS "
       "output (Exhibit 18).")
bullet("RegMap released as a reusable, openly-licensed model and reused as a component across the "
       "platform (Exhibit 11A).")
para(
    "This exhibit should be read alongside Exhibits 11–18, which document the "
    "individual prototypes, their live validation, and their consolidation into an operational "
    "platform, and alongside my publication record and personal statement.")

para("")
para("Date: " + AS_OF)
para("Prepared by: Samuel Gbli Tetteh")

out = r"c:\Users\User\ai-cybersecurity-portfolio\exhibits\Exhibit 15 Research Progress and Body of Work.docx"
doc.save(out)
print("saved:", out)
print("paragraphs:", len(doc.paragraphs), "| tables:", len(doc.tables))
