"""Generate "Exhibit 12A - Hybrid Identity Anomaly Detector Released as a Standalone System.docx"."""
from docx import Document
from docx.shared import Pt

RUN_DATE = "July 17, 2026"
REPO = "https://github.com/samuelgtetteh/ai-cybersecurity-portfolio"

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


def para(t, bold=False, italic=False):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = bold; r.italic = italic; return p


def bullet(t):
    doc.add_paragraph(t, style="List Bullet")


def table(headers, rows):
    tb = doc.add_table(rows=1, cols=len(headers)); tb.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        tb.rows[0].cells[j].paragraphs[0].add_run(h).bold = True
    for row in rows:
        c = tb.add_row().cells
        for j, v in enumerate(row):
            c[j].text = str(v)
    return tb


t = doc.add_paragraph()
r = t.add_run("Exhibit 12A: Hybrid Identity Anomaly Detector — Released as a Standalone System")
r.bold = True; r.font.size = Pt(15)
para("Petitioner: Samuel Gbli Tetteh", bold=True)
para("Addendum to Exhibit 12 — documenting that the identity anomaly detector has been packaged and "
     "publicly released as a reusable, standalone tool.", italic=True)

doc.add_heading("1. Packaged and published for public use", level=1)
para("The Hybrid Identity Anomaly Detector (an Isolation Forest that flags anomalous authentication "
     "events — credential compromise, lateral movement, insider threats) has been released as a "
     "self-contained tool anyone can run, under the permissive Apache-2.0 license.")
table(["Channel", "Location"],
      [["Hugging Face", "huggingface.co/stetteh/identity-anomaly"],
       ["GitHub Release", "v0.1-identity (model + scaler + inference wrapper, runs offline)"],
       ["Docker image", "ghcr.io/samuelgtetteh/identity-anomaly (serving API: POST /score)"]])

doc.add_heading("2. What the release contains", level=1)
bullet("The trained model + fitted scaler, plus a scoring wrapper that reproduces the exact "
       "9-feature engineering and 1-hour rolling-window aggregates used at serving time.")
bullet("A FastAPI serving image (POST /score) and a runnable example.")
bullet("A model card stating intended use and limitations — assistive detection for a SOC to "
       "triage, not an automated verdict.")

doc.add_heading("3. Significance", level=1)
para("The detector is not a private prototype: it is openly available, reusable, and also operates "
     "as a component of the real-time decisioning platform (Exhibits 16–19). Publishing it as a "
     "standalone system that U.S. organizations can adopt at no cost reinforces both that the "
     "petitioner is well positioned to advance the endeavor and that the work is disseminated for "
     "public benefit.")
p = doc.add_paragraph(); p.add_run(REPO).bold = True
para(""); para("Date: " + RUN_DATE); para("Prepared by: Samuel Gbli Tetteh")

out = r"c:\Users\User\ai-cybersecurity-portfolio\exhibits\Exhibit 12A Identity Detector Standalone Release.docx"
doc.save(out); print("saved:", out, "| paras", len(doc.paragraphs), "| tables", len(doc.tables))
