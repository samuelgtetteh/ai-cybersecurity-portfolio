"""
Build downloadable SCAN-FINDINGS reports (JSON / XLSX / DOCX) from an analyze() result.

This is the raw technical assessment (assets -> services -> versions -> CVEs with KEV/EPSS/risk +
the NIST category), distinct from the compliance report_export (which is control-focused). Both
are offered so the user can download the findings and independently CHOOSE whether to continue to
compliance mapping.

    build_all(outdir, analysis) -> {"json": "findings.json", "xlsx": "findings.xlsx", "docx": "findings.docx"}

`analysis` is the object returned by advisor_api._analyze / securescan.analyze.analyze:
  {"scan_report": {...}, "hosts": [{ip, categories, ports:[{port,service,product,version,cpe,
   category, cves:[{cve_id, cvss_score, in_kev, epss, risk, summary}]}]}], "cve_total", "kev_count",
   "host_max_risk", "categories": [...]}
"""
import json
from datetime import datetime, timezone
from pathlib import Path


def _now():
    return datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")


def _rows(analysis):
    """Flatten to one row per (host, port, cve) — a port with no CVEs still yields one row."""
    rows = []
    for h in analysis.get("hosts", []):
        ip = h.get("ip")
        for p in h.get("ports", []):
            base = {"host": ip, "port": p.get("port"), "service": p.get("service", ""),
                    "product": p.get("product", ""), "version": p.get("version", ""),
                    "cpe": p.get("cpe", ""), "category": p.get("category", "")}
            cves = p.get("cves") or []
            if not cves:
                rows.append({**base, "cve": "", "cvss": None, "kev": "", "epss": None,
                             "risk": None, "summary": ""})
                continue
            for c in cves:
                rows.append({**base, "cve": c.get("cve_id", ""), "cvss": c.get("cvss_score"),
                             "kev": "yes" if c.get("in_kev") else "", "epss": c.get("epss"),
                             "risk": c.get("risk"),
                             "summary": (c.get("summary") or c.get("description") or "")[:300]})
    return rows


_COLS = ["host", "port", "service", "product", "version", "cve", "cvss", "kev", "epss", "risk",
         "category", "cpe", "summary"]


def _summary(analysis):
    hosts = analysis.get("hosts", [])
    return {"captured": _now(), "cidr": (analysis.get("scan_report") or {}).get("cidr"),
            "assets": len(hosts), "open_services": sum(len(h.get("ports", [])) for h in hosts),
            "cve_total": analysis.get("cve_total", 0), "kev_count": analysis.get("kev_count", 0),
            "host_max_risk": analysis.get("host_max_risk", 0),
            "categories": analysis.get("categories", [])}


def build_json(path, analysis):
    Path(path).write_text(json.dumps(
        {"summary": _summary(analysis), "rows": _rows(analysis), "analysis": analysis},
        default=str, indent=2), encoding="utf-8")


def build_xlsx(path, analysis):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    wb = Workbook()
    s = _summary(analysis)
    # Summary sheet
    ov = wb.active; ov.title = "Summary"
    ov["A1"] = "RegMap — Scan Findings"; ov["A1"].font = Font(bold=True, size=14)
    meta = [("Captured", s["captured"]), ("Target", s["cidr"] or "single host"),
            ("Assets", s["assets"]), ("Open services", s["open_services"]),
            ("Weaknesses (CVEs)", s["cve_total"]), ("Known-exploited (KEV)", s["kev_count"]),
            ("Max host risk", s["host_max_risk"]), ("Categories", ", ".join(s["categories"]))]
    for i, (k, v) in enumerate(meta, start=3):
        ov.cell(row=i, column=1, value=k).font = Font(bold=True)
        ov.cell(row=i, column=2, value=v)
    ov.column_dimensions["A"].width = 22; ov.column_dimensions["B"].width = 60
    # Findings sheet
    fs = wb.create_sheet("Findings")
    hdr = Font(bold=True, color="FFFFFF"); fill = PatternFill("solid", fgColor="1668C9")
    for j, c in enumerate(_COLS, start=1):
        cell = fs.cell(row=1, column=j, value=c.upper()); cell.font = hdr; cell.fill = fill
    for i, r in enumerate(_rows(analysis), start=2):
        for j, c in enumerate(_COLS, start=1):
            fs.cell(row=i, column=j, value=r.get(c))
    widths = {"host": 15, "product": 26, "cve": 16, "summary": 60, "cpe": 26}
    for j, c in enumerate(_COLS, start=1):
        fs.column_dimensions[fs.cell(row=1, column=j).column_letter].width = widths.get(c, 12)
    fs.freeze_panes = "A2"; fs.auto_filter.ref = fs.dimensions
    wb.save(path)


def build_docx(path, analysis):
    from docx import Document
    from docx.shared import Pt, RGBColor
    s = _summary(analysis)
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10.5)
    t = doc.add_paragraph(); r = t.add_run("RegMap — Scan Findings")
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x44)
    doc.add_paragraph("Captured %s · target %s" % (s["captured"], s["cidr"] or "single host")).italic = True
    doc.add_heading("Summary", level=1)
    tb = doc.add_table(rows=0, cols=2); tb.style = "Light Grid Accent 1"
    for k, v in [("Assets", s["assets"]), ("Open services", s["open_services"]),
                 ("Weaknesses (CVEs)", s["cve_total"]), ("Known-exploited (KEV)", s["kev_count"]),
                 ("Max host risk", s["host_max_risk"]), ("Categories", ", ".join(s["categories"]))]:
        c = tb.add_row().cells; c[0].paragraphs[0].add_run(str(k)).bold = True; c[1].text = str(v)
    for h in analysis.get("hosts", []):
        doc.add_heading("%s" % h.get("ip"), level=2)
        if h.get("categories"):
            doc.add_paragraph("Categories: " + ", ".join(h["categories"]))
        pt = doc.add_table(rows=1, cols=5); pt.style = "Light Grid Accent 1"
        for j, hd in enumerate(["Port/Service", "Product", "CVE", "CVSS / KEV / EPSS", "Summary"]):
            pt.rows[0].cells[j].paragraphs[0].add_run(hd).bold = True
        for p in h.get("ports", []):
            svc = "%s/%s" % (p.get("port"), p.get("service") or "")
            prod = " ".join(x for x in [p.get("product"), p.get("version")] if x) or "—"
            cves = p.get("cves") or []
            if not cves:
                row = pt.add_row().cells
                row[0].text = svc; row[1].text = prod; row[2].text = "—"; row[3].text = "—"; row[4].text = "no known CVEs"
            for c in cves:
                row = pt.add_row().cells
                row[0].text = svc; row[1].text = prod; row[2].text = c.get("cve_id", "")
                kev = "KEV " if c.get("in_kev") else ""
                epss = ("%d%%" % round((c.get("epss") or 0) * 100))
                row[3].text = "%s  %s%s" % (c.get("cvss_score"), kev, epss)
                row[4].text = (c.get("summary") or c.get("description") or "")[:400]
    doc.save(path)


def build_all(outdir, analysis):
    outdir = Path(outdir)
    files = {}
    build_json(outdir / "findings.json", analysis); files["json"] = "findings.json"
    build_xlsx(outdir / "findings.xlsx", analysis); files["xlsx"] = "findings.xlsx"
    build_docx(outdir / "findings.docx", analysis); files["docx"] = "findings.docx"
    return files
